# -*- coding: utf-8 -*-
import sys
import io
import re
from docx import Document
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc_path = r'C:\Users\KiKi\Desktop\ai\AI_assign\output\report.docx'

try:
    doc = Document(doc_path)
    print('파일 열기 성공')
except Exception as e:
    print(f'파일 열기 실패: {e}')
    sys.exit(1)

# ─────────────────────────────────────────────
# 전체 텍스트 수집 (본문 + 표)
# ─────────────────────────────────────────────
all_para_texts = [para.text for para in doc.paragraphs]
table_texts = []
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                table_texts.append(para.text)

full_text = '\n'.join(all_para_texts + table_texts)

# ─────────────────────────────────────────────
# [1] 그림 번호 연속성
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('[1] 그림 번호 연속성 검증')
print('=' * 60)
figure_refs = re.findall(r'\[그림\s*(\d+)\]', full_text)
figure_nums_unique = sorted(set(int(n) for n in figure_refs))
count_map = {}
for n in figure_refs:
    count_map[int(n)] = count_map.get(int(n), 0) + 1

print(f'발견된 그림 번호: {figure_nums_unique}')
missing = [i for i in range(1, 16) if i not in figure_nums_unique]
extra   = [n for n in figure_nums_unique if n < 1 or n > 15]
duplicates = {k: v for k, v in count_map.items() if v > 1}

print(f'누락된 번호 (1~15 중): {missing if missing else "없음"}')
print(f'범위 외 번호: {extra if extra else "없음"}')
print(f'중복 등장 번호: {duplicates if duplicates else "없음"}')
result1 = not missing and not extra
print(f'결과: {"PASS" if result1 else "FAIL"}')

# ─────────────────────────────────────────────
# [2] 이미지 개수
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('[2] 이미지 개수 검증')
print('=' * 60)
blip_count = 0
for para in doc.paragraphs:
    blip_count += len(para._element.findall('.//' + qn('a:blip')))
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                blip_count += len(para._element.findall('.//' + qn('a:blip')))
print(f'총 삽입 이미지 수: {blip_count}개')
result2 = blip_count == 15
print(f'결과: {"PASS" if result2 else "FAIL"} (기대값: 15개)')

# ─────────────────────────────────────────────
# [3] 표 개수
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('[3] 표 개수 검증')
print('=' * 60)
table_count = len(doc.tables)
print(f'총 표 개수: {table_count}개')

# ─────────────────────────────────────────────
# [4] 섹션 구조
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('[4] 섹션 구조 검증 (1절~5절 헤딩)')
print('=' * 60)

print('--- Heading 스타일 단락 목록 ---')
heading_paras = []
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        heading_paras.append((para.style.name, para.text))
        print(f'  [{para.style.name}] {para.text[:80]}')

if not heading_paras:
    print('  (Heading 스타일 단락 없음)')

print()
print('--- "절" 포함 단락 목록 ---')
section_found = []
for para in doc.paragraphs:
    t = para.text
    if re.search(r'[1-5]\s*절', t):
        section_found.append(t.strip())
        print(f'  [{para.style.name}] {t[:100]}')

if not section_found:
    print('  (없음 — 다른 패턴으로 재탐색)')
    # "제N절" 또는 "N." 숫자 형식
    for para in doc.paragraphs:
        t = para.text
        if re.search(r'제\s*[1-5]\s*절', t) or re.match(r'^[1-5]\.\s+\S', t):
            section_found.append(t.strip())
            print(f'  [{para.style.name}] {t[:100]}')

expected_labels = ['1절', '2절', '3절', '4절', '5절']
missing_sections = [lbl for lbl in expected_labels if not any(lbl in s for s in section_found)]
print(f'\n누락된 절 헤딩: {missing_sections if missing_sections else "없음"}')
result4 = not missing_sections
print(f'결과: {"PASS" if result4 else "FAIL"}')

# ─────────────────────────────────────────────
# [5] 핵심 텍스트 존재 여부
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('[5] 핵심 텍스트 존재 여부')
print('=' * 60)
keywords = [
    'ARIMA 채택',
    '2026-05-25',
    'SimpleRNN',
    '수준 상관',
    '계절 분해',
    '시기별·지역별 누적',
]
all_kw_pass = True
for kw in keywords:
    found = kw in full_text
    status = 'PASS' if found else 'FAIL'
    if not found:
        all_kw_pass = False
    print(f'  [{status}] "{kw}"')
result5 = all_kw_pass
print(f'결과: {"PASS" if result5 else "FAIL"}')

# ─────────────────────────────────────────────
# 최종 요약
# ─────────────────────────────────────────────
print()
print('=' * 60)
print('최종 요약')
print('=' * 60)
results = {
    '1. 그림 번호 연속성': result1,
    '2. 이미지 개수(=15)': result2,
    '3. 표 개수': None,       # 단순 정보 항목
    '4. 섹션 구조': result4,
    '5. 핵심 텍스트': result5,
}
for label, res in results.items():
    if res is None:
        print(f'  [INFO] {label}: {table_count}개')
    else:
        print(f'  [{"PASS" if res else "FAIL"}] {label}')
